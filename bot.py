import os
from collections import defaultdict, deque

from telegram import Update, ReplyKeyboardMarkup, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import (
    ApplicationBuilder,
    MessageHandler,
    CommandHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters
)

from docx import Document
import pdfplumber
import openpyxl
import csv
import markdown2
from PIL import Image
import io

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import urllib.request

# Download a Unicode font (DejaVu) once at startup if not present
FONT_PATH = "/tmp/DejaVuSans.ttf"
FONT_URL  = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"

def ensure_font():
    if not os.path.exists(FONT_PATH):
        urllib.request.urlretrieve(FONT_URL, FONT_PATH)
    try:
        pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))
    except Exception:
        pass  # already registered

ensure_font()


TOKEN = os.getenv("TOKEN")
ADMIN_ID = 366339367  # <-- ВСТАВЬ СВОЙ TELEGRAM ID


# ---------------- STORAGE ----------------
user_stats = defaultdict(int)
user_history = defaultdict(lambda: deque(maxlen=5))
user_names = {}          # user_id -> "Name (@username)"
pending_files = {}       # user_id -> {"path": ..., "name": ..., "ext": ...}


# ================================================================
#  PDF HELPER  (reportlab, full Unicode via DejaVu font)
# ================================================================
def _lines_to_pdf(lines, dst):
    """Write a list of text lines to a PDF with proper Unicode support."""
    doc = SimpleDocTemplate(
        dst,
        pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm,  bottomMargin=20*mm,
    )
    style = ParagraphStyle(
        "body",
        fontName="DejaVu",
        fontSize=11,
        leading=16,
        wordWrap="CJK",   # handles all scripts
    )
    story = []
    for line in lines:
        text = line.strip() if isinstance(line, str) else line
        # Escape XML special chars for reportlab
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if text:
            story.append(Paragraph(text, style))
        else:
            story.append(Spacer(1, 6))
    if not story:
        story.append(Paragraph("(empty)", style))
    doc.build(story)


# ================================================================
#  CONVERTERS
# ================================================================

# --- DOCX ---
def docx_to_txt(src, dst):
    doc = Document(src)
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n".join(p.text for p in doc.paragraphs))

def docx_to_md(src, dst):
    doc = Document(src)
    lines = []
    for p in doc.paragraphs:
        style = p.style.name.lower()
        if "heading 1" in style:
            lines.append(f"# {p.text}")
        elif "heading 2" in style:
            lines.append(f"## {p.text}")
        elif "heading 3" in style:
            lines.append(f"### {p.text}")
        else:
            lines.append(p.text)
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n\n".join(lines))

def docx_to_pdf(src, dst):
    doc = Document(src)
    lines = [p.text for p in doc.paragraphs]
    _lines_to_pdf(lines, dst)


# --- TXT ---
def txt_to_docx(src, dst):
    doc = Document()
    with open(src, "r", encoding="utf-8") as f:
        for line in f:
            doc.add_paragraph(line.strip())
    doc.save(dst)

def txt_to_md(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        content = f.read()
    with open(dst, "w", encoding="utf-8") as f:
        f.write(content)   # plain text is valid markdown

def txt_to_pdf(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        lines = f.readlines()
    _lines_to_pdf(lines, dst)


# --- PDF ---
def pdf_to_txt(src, dst):
    with pdfplumber.open(src) as pdf:
        pages = []
        for page in pdf.pages:
            # extract_text with layout=True preserves spacing & columns
            text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
            if text:
                pages.append(text)
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n\n--- Page Break ---\n\n".join(pages))

def pdf_to_docx(src, dst):
    with pdfplumber.open(src) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
            if text:
                pages.append(text)
    doc = Document()
    for i, page_text in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        for line in page_text.splitlines():
            doc.add_paragraph(line)
    doc.save(dst)

def pdf_to_md(src, dst):
    with pdfplumber.open(src) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
            if text:
                pages.append(text)
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n\n---\n\n".join(pages))


# --- XLSX / CSV ---
def xlsx_to_csv(src, dst):
    wb = openpyxl.load_workbook(src)
    ws = wb.active
    with open(dst, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow([v if v is not None else "" for v in row])

def csv_to_xlsx(src, dst):
    wb = openpyxl.Workbook()
    ws = wb.active
    with open(src, "r", encoding="utf-8") as f:
        for row in csv.reader(f):
            ws.append(row)
    wb.save(dst)

def xlsx_to_txt(src, dst):
    wb = openpyxl.load_workbook(src)
    ws = wb.active
    lines = []
    for row in ws.iter_rows(values_only=True):
        lines.append("\t".join(str(v) if v is not None else "" for v in row))
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

def csv_to_txt(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        content = f.read()
    with open(dst, "w", encoding="utf-8") as f:
        f.write(content)


# --- Markdown ---
def md_to_txt(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        content = f.read()
    # strip simple markdown symbols
    import re
    plain = re.sub(r"[#*_`>\-\[\]!]", "", content)
    with open(dst, "w", encoding="utf-8") as f:
        f.write(plain)

def md_to_docx(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        lines = f.readlines()
    doc = Document()
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            doc.add_paragraph(stripped)
    doc.save(dst)

def md_to_pdf(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        lines = f.readlines()
    _lines_to_pdf(lines, dst)


# --- Images ---
def img_to_jpg(src, dst):
    Image.open(src).convert("RGB").save(dst, "JPEG")

def img_to_png(src, dst):
    Image.open(src).save(dst, "PNG")

def img_to_pdf(src, dst):
    image = Image.open(src).convert("RGB")
    image.save(dst, "PDF", resolution=100.0)

def img_to_txt(src, dst):
    """Best-effort: writes a note since OCR (tesseract) may not be installed."""
    try:
        import pytesseract
        text = pytesseract.image_to_string(Image.open(src))
    except Exception:
        text = "[OCR not available – install pytesseract + tesseract to extract text from images]"
    with open(dst, "w", encoding="utf-8") as f:
        f.write(text)


# ================================================================
#  CONVERSION ROUTING TABLE
#  { source_ext: { target_ext: converter_fn } }
# ================================================================
CONVERTERS = {
    ".docx": {
        ".txt":  docx_to_txt,
        ".md":   docx_to_md,
        ".pdf":  docx_to_pdf,
    },
    ".txt": {
        ".docx": txt_to_docx,
        ".md":   txt_to_md,
        ".pdf":  txt_to_pdf,
    },
    ".pdf": {
        ".txt":  pdf_to_txt,
        ".docx": pdf_to_docx,
        ".md":   pdf_to_md,
    },
    ".xlsx": {
        ".csv":  xlsx_to_csv,
        ".txt":  xlsx_to_txt,
    },
    ".csv": {
        ".xlsx": csv_to_xlsx,
        ".txt":  csv_to_txt,
    },
    ".md": {
        ".txt":  md_to_txt,
        ".docx": md_to_docx,
        ".pdf":  md_to_pdf,
    },
    ".jpg": {
        ".png":  img_to_png,
        ".pdf":  img_to_pdf,
        ".txt":  img_to_txt,
    },
    ".jpeg": {
        ".png":  img_to_png,
        ".pdf":  img_to_pdf,
        ".txt":  img_to_txt,
    },
    ".png": {
        ".jpg":  img_to_jpg,
        ".pdf":  img_to_pdf,
        ".txt":  img_to_txt,
    },
}

SUPPORTED_EXTS = set(CONVERTERS.keys())

# Human-readable labels
EXT_LABELS = {
    ".docx": "DOCX",
    ".txt":  "TXT",
    ".pdf":  "PDF",
    ".md":   "Markdown",
    ".xlsx": "XLSX",
    ".csv":  "CSV",
    ".jpg":  "JPG",
    ".jpeg": "JPG",
    ".png":  "PNG",
}


# ================================================================
#  MENU
# ================================================================
menu = ReplyKeyboardMarkup(
    [
        ["📎 Отправить файл"],
        ["📊 Статистика", "🧾 История"],
        ["ℹ️ Помощь"]
    ],
    resize_keyboard=True
)


# ================================================================
#  HANDLERS
# ================================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    supported = ", ".join(sorted({EXT_LABELS[e] for e in SUPPORTED_EXTS}))
    await update.message.reply_text(
        f"👋 Добро пожаловать в PRO-конвертер!\n\n"
        f"🛠 Поддерживаемые форматы:\n{supported}\n\n"
        f"📎 Просто отправь файл — бот предложит варианты конвертации 👇",
        reply_markup=menu
    )


async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    lines = []
    for src, targets in CONVERTERS.items():
        src_label = EXT_LABELS.get(src, src)
        tgt_labels = ", ".join(EXT_LABELS.get(t, t) for t in targets)
        lines.append(f"• {src_label} → {tgt_labels}")
    await update.message.reply_text(
        "ℹ️ Отправь файл и выбери нужный формат.\n\n"
        "Доступные конвертации:\n" + "\n".join(lines)
    )


async def stats(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    count = user_stats[user_id]
    user = user_names.get(user_id, "Ты")
    await update.message.reply_text(f"📊 Статистика\n\n👤 {user}\n📎 Файлов: {count}")


async def history(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    hist = user_history[user_id]
    if not hist:
        await update.message.reply_text("🧾 История пуста")
        return
    text = "🧾 Последние файлы:\n\n" + "".join(f"📄 {f}\n" for f in hist)
    await update.message.reply_text(text)


# ---------------- ADMIN ----------------
def is_admin(user_id):
    return user_id == ADMIN_ID

async def admin_panel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.message.from_user.id):
        await update.message.reply_text("⛔ Нет доступа"); return
    await update.message.reply_text(
        "👑 АДМИН-ПАНЕЛЬ\n\n/users - пользователи\n/allstats - статистика\n/logs - действия"
    )

async def users_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.message.from_user.id):
        await update.message.reply_text("⛔ Нет доступа"); return
    if not user_names:
        await update.message.reply_text("Нет пользователей"); return
    text = "👥 Пользователи:\n\n" + "".join(f"{n} | ID: {uid}\n" for uid, n in user_names.items())
    await update.message.reply_text(text)

async def all_stats(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.message.from_user.id):
        await update.message.reply_text("⛔ Нет доступа"); return
    total = sum(user_stats.values())
    text = f"📊 Общая статистика\n\nВсего файлов: {total}\n\n"
    text += "".join(f"{user_names.get(uid,'Unknown')}: {c}\n" for uid, c in user_stats.items())
    await update.message.reply_text(text)

async def logs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.message.from_user.id):
        await update.message.reply_text("⛔ Нет доступа"); return
    text = "🧾 Последние действия:\n\n"
    for uid, files in user_history.items():
        text += f"{user_names.get(uid,'Unknown')}:\n" + "".join(f"  - {f}\n" for f in files) + "\n"
    await update.message.reply_text(text)


# ---------------- FILE HANDLER ----------------
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        file = update.message.document
        if not file:
            return

        user_obj   = update.message.from_user
        user_id    = user_obj.id
        name       = user_obj.first_name
        username   = user_obj.username
        user_names[user_id] = f"{name} (@{username})" if username else f"{name} (@нет)"

        file_name = file.file_name
        _, ext = os.path.splitext(file_name.lower())

        if ext not in SUPPORTED_EXTS:
            supported = ", ".join(sorted({EXT_LABELS[e] for e in SUPPORTED_EXTS}))
            await update.message.reply_text(
                f"❌ Формат не поддерживается.\n\nПоддерживаются: {supported}"
            )
            return

        # Download
        new_file   = await file.get_file()
        input_path = f"input_{user_id}_{file_name}"
        await new_file.download_to_drive(input_path)

        # Store pending file info
        pending_files[user_id] = {"path": input_path, "name": file_name, "ext": ext}

        # Build inline keyboard with target options
        targets = CONVERTERS[ext]
        buttons = [
            InlineKeyboardButton(
                f"→ {EXT_LABELS.get(t, t.upper())}",
                callback_data=f"convert:{t}"
            )
            for t in targets
        ]
        # Two buttons per row
        keyboard = [buttons[i:i+2] for i in range(0, len(buttons), 2)]

        await update.message.reply_text(
            f"📄 Файл получен: `{file_name}`\n\nВыбери формат конвертации:",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode="Markdown"
        )

    except Exception as e:
        await update.message.reply_text(f"⚠️ Ошибка: {str(e)}")


# ---------------- CALLBACK: CONVERSION CHOICE ----------------
async def conversion_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    data    = query.data  # "convert:.ext"

    if not data.startswith("convert:"):
        return

    target_ext = data[len("convert:"):]

    pending = pending_files.get(user_id)
    if not pending:
        await query.edit_message_text("⚠️ Файл не найден. Отправь его заново.")
        return

    input_path = pending["path"]
    file_name  = pending["name"]
    src_ext    = pending["ext"]

    base_name   = os.path.splitext(file_name)[0]
    output_name = base_name + target_ext
    output_path = f"output_{user_id}_{output_name}"

    await query.edit_message_text(f"⚡ Конвертирую в {EXT_LABELS.get(target_ext, target_ext)}...")

    try:
        converter = CONVERTERS[src_ext][target_ext]
        converter(input_path, output_path)

        with open(output_path, "rb") as f:
            await context.bot.send_document(
                chat_id=user_id,
                document=f,
                filename=output_name,
                caption=f"✅ Готово! {EXT_LABELS.get(src_ext,src_ext)} → {EXT_LABELS.get(target_ext,target_ext)}"
            )

        # Notify admin
        try:
            user_info = pending_files[user_id]
            caption = (
                f"📥 Новый файл\n\n"
                f"👤 {query.from_user.first_name}\n"
                f"🔗 @{query.from_user.username or 'нет'}\n"
                f"🆔 {user_id}\n"
                f"📄 {file_name} → {output_name}"
            )
            with open(input_path, "rb") as f:
                await context.bot.send_document(chat_id=ADMIN_ID, document=f, caption=caption)
        except Exception as e:
            print("Ошибка отправки админу:", e)

        user_stats[user_id]   += 1
        user_history[user_id].append(f"{file_name} → {output_name}")

    except Exception as e:
        await context.bot.send_message(chat_id=user_id, text=f"⚠️ Ошибка конвертации: {str(e)}")

    finally:
        for p in (input_path, output_path):
            try:
                os.remove(p)
            except Exception:
                pass
        pending_files.pop(user_id, None)


# ---------------- TEXT / BUTTON HANDLER ----------------
async def text_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    if text == "📎 Отправить файл":
        await update.message.reply_text("📎 Отправь файл — я скажу, в что можно конвертировать")
    elif text == "📊 Статистика":
        await stats(update, context)
    elif text == "🧾 История":
        await history(update, context)
    elif text == "ℹ️ Помощь":
        await help_cmd(update, context)


# ================================================================
#  MAIN
# ================================================================
def main():
    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start",    start))
    app.add_handler(CommandHandler("help",     help_cmd))
    app.add_handler(CommandHandler("admin",    admin_panel))
    app.add_handler(CommandHandler("users",    users_list))
    app.add_handler(CommandHandler("allstats", all_stats))
    app.add_handler(CommandHandler("logs",     logs))

    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, text_handler))
    app.add_handler(CallbackQueryHandler(conversion_callback, pattern=r"^convert:"))

    app.run_polling()


if __name__ == "__main__":
    main()
